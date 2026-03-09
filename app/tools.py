# app/tools.py
"""
수입 비용 계산 에이전트를 위한 도구 정의
비동기 지원 및 명확한 리턴 타입 제공
"""
import os
import re
import asyncio
import requests
import pandas as pd
import numpy as np
from typing import List, Dict, Any, Optional
from dotenv import load_dotenv
from langchain_core.tools import tool
from langchain_community.vectorstores import FAISS
from app.models import get_embedding_model
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from docx import Document as DocxDocument

# 환경 변수 로드
load_dotenv()

# --- CSV 파일을 미리 로드하여 메모리에 저장 ---
TARIFF_DF = None
try:
    TARIFF_DF = pd.read_csv("./data/tariff_by_hs.csv", encoding="cp949", dtype={"세번": str, "잠정세율 - B": str})
    print("tariff_by_hs.csv 파일을 성공적으로 로드했습니다.")
except FileNotFoundError:
    print("경고: ./data/tariff_by_hs.csv 파일을 찾을 수 없습니다.")


# --- 검색 스토어 (Lazy Loading) ---
_vector_store = None       # 통합 인덱스 (PDF + CSV)
_pdf_store    = None       # PDF-only 인덱스 (balanced search용)
_hs_code_search_call_count = 0
_tariff_search_call_count  = 0
_query_expander = None

# balanced search 설정 (RAG 실험 최적값: PDF 쿼터=3)
PDF_QUOTA        = 3
COMBINED_VS_PATH = "./vector_store/faiss_index"
PDF_ONLY_VS_PATH = "./vector_store/faiss_index_pdf"

# 검색 쿼리 확장 (LLM 기반) 사용 여부
# ReAct 에이전트가 직접 쿼리를 제어하므로 expansion 비활성화
# (활성화 시 ReAct가 자신의 쿼리가 내부에서 변환되는 사실을 모르고 중복 호출 발생)
USE_LLM_QUERY_EXPANSION = False


def reset_hs_code_search_limit() -> None:
    """HS 코드/관세 검색 도구 호출 카운터를 리셋 (에이전트 1회 실행 전 호출)."""
    global _hs_code_search_call_count, _tariff_search_call_count
    _hs_code_search_call_count = 0
    _tariff_search_call_count = 0


def _load_stores() -> None:
    """통합 인덱스 + PDF-only 인덱스 Lazy Loading."""
    global _vector_store, _pdf_store
    if _vector_store is not None:
        return

    if not os.path.exists(COMBINED_VS_PATH):
        raise FileNotFoundError(
            f"Vector store not found: {COMBINED_VS_PATH}\n"
            "Please run 'python run_preprocessing.py' first."
        )

    embedding_model = get_embedding_model()
    _vector_store = FAISS.load_local(
        COMBINED_VS_PATH, embedding_model, allow_dangerous_deserialization=True
    )

    if os.path.exists(PDF_ONLY_VS_PATH):
        _pdf_store = FAISS.load_local(
            PDF_ONLY_VS_PATH, embedding_model, allow_dangerous_deserialization=True
        )
        print(f"[Tools] 듀얼 인덱스 로드 완료 (PDF 쿼터={PDF_QUOTA})")
    else:
        print("[Tools] PDF-only 인덱스 없음 — 단일 인덱스 검색으로 동작합니다.")


def _balanced_search(query: str, k: int = 5) -> List:
    """
    통합 인덱스에서 (k - PDF_QUOTA)개 + PDF-only 인덱스에서 PDF_QUOTA개 검색.
    PDF-only 인덱스가 없으면 통합 인덱스 단독 검색으로 fallback.
    """
    _load_stores()
    if _pdf_store is None:
        return _vector_store.similarity_search(query, k=k)

    csv_k = max(0, k - PDF_QUOTA)
    pdf_k = min(k, PDF_QUOTA)
    csv_docs = _vector_store.similarity_search(query, k=csv_k) if csv_k > 0 else []
    pdf_docs = _pdf_store.similarity_search(query, k=pdf_k)
    return csv_docs + pdf_docs


def get_retriever():
    """저장된 FAISS 인덱스로부터 Retriever를 반환 (내부 호환용)."""
    _load_stores()
    return _vector_store.as_retriever(search_kwargs={"k": 5})


def get_query_expander():
    """검색 쿼리 확장용 LLM (Lazy Loading)"""
    global _query_expander
    if _query_expander is not None:
        return _query_expander
    from langchain_openai import ChatOpenAI
    _query_expander = ChatOpenAI(model="gpt-4o-mini", temperature=0)
    return _query_expander


def expand_hs_search_query(query: str) -> str:
    """
    HS 코드 검색 정확도를 높이기 위해 검색 문장을 확장합니다.
    - 재질/용도/형태/구성 요소를 포함한 한두 문장 생성
    - 과장/허위 정보 금지
    """
    if not USE_LLM_QUERY_EXPANSION:
        return query
    try:
        llm = get_query_expander()
        prompt = (
            "다음 상품명/설명을 HS 코드 검색에 유리하도록 확장하세요.\n"
            "- 재질/용도/형태/구성 요소 중심\n"
            "- 과장/허위 금지\n"
            "- 1~2문장, 120자 이내\n"
            f"상품: {query}\n"
            "확장 검색문장:"
        )
        resp = llm.invoke(prompt)
        expanded = (resp.content or "").strip()
        # 너무 길면 축약
        if len(expanded) > 120:
            expanded = expanded[:120].strip()
        return expanded or query
    except Exception:
        return query


def _score_doc_for_snippet(text: str, hs_code: Optional[str], item_name: Optional[str]) -> int:
    """문서에 대한 스니펫 우선순위 점수 반환 (높을수록 우선)."""
    score = 0
    hs_digits = re.sub(r"\D", "", hs_code) if hs_code else ""
    tokens = [t for t in re.split(r"\s+", item_name) if len(t) > 1] if item_name else []

    if hs_digits and hs_digits in re.sub(r"\D", "", text):
        score += 10  # HS 코드 포함 문서 최우선
    if tokens:
        matched = sum(1 for tok in tokens if tok in text)
        score += matched * 2
    # CSV 단일 행 (너무 짧으면 페널티)
    if len(text.strip()) < 30:
        score -= 5
    return score


def _extract_snippet_from_doc(text: str, hs_code: Optional[str], item_name: Optional[str],
                               max_len: int = 300) -> str:
    """
    문서에서 HS 코드·물품명 주변 맥락을 포함한 스니펫 추출.
    - HS 코드가 포함된 줄을 찾아 그 앞뒤 문장을 포함해 반환
    - 없으면 전체 텍스트를 max_len 안에서 반환
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return text.strip()[:max_len]

    hs_digits = re.sub(r"\D", "", hs_code) if hs_code else ""
    tokens = [t for t in re.split(r"\s+", item_name) if len(t) > 1] if item_name else []

    # HS 코드 포함 줄 인덱스 탐색
    anchor_idx = None
    if hs_digits:
        for i, ln in enumerate(lines):
            if hs_digits in re.sub(r"\D", "", ln):
                anchor_idx = i
                break

    # 물품명 토큰 포함 줄 (HS 코드 없을 때 차선)
    if anchor_idx is None and tokens:
        for i, ln in enumerate(lines):
            if any(tok in ln for tok in tokens):
                anchor_idx = i
                break

    if anchor_idx is None:
        anchor_idx = 0

    # anchor 줄 기준 앞뒤 1줄 포함, max_len 이내
    start = max(0, anchor_idx - 1)
    end   = min(len(lines), anchor_idx + 3)
    snippet = " ".join(lines[start:end])
    return snippet[:max_len]


def get_rag_snippet_for_candidate(item_name: str, hs_code: str, llm_snippet: Optional[str] = None) -> str:
    """
    RAG 검색 결과에서 관세청 DB 원문 스니펫을 반환합니다.
    - PDF 청크를 우선 탐색 (HS 코드·물품명 맥락이 풍부)
    - CSV 단일 행은 내용이 빈약하므로 낮은 우선순위
    - LLM이 제시한 인용이 실제 DB에 포함되면 보강 재료로 활용
    """
    try:
        _load_stores()
        docs: List = []

        # PDF-only 인덱스 우선 (분류 기준 원문이 풍부)
        if _pdf_store is not None:
            docs = _pdf_store.similarity_search(f"{item_name} {hs_code}", k=5)
        # PDF-only 없거나 결과 부족하면 통합 인덱스 보충
        if len(docs) < 3 and _vector_store is not None:
            docs += _vector_store.similarity_search(f"{item_name} {hs_code}", k=5)

        if not docs:
            return llm_snippet or ""

        # 스코어 기반 정렬 → 가장 관련성 높은 문서에서 스니펫 추출
        scored = sorted(
            docs,
            key=lambda d: _score_doc_for_snippet(d.page_content, hs_code, item_name),
            reverse=True,
        )
        for doc in scored:
            snippet = _extract_snippet_from_doc(doc.page_content, hs_code, item_name)
            if snippet and len(snippet) > 10:
                return snippet

    except Exception:
        pass

    return llm_snippet or ""


# ===== HS Code & Tax Finder 도구 =====

@tool("hs_code_search")
def hs_code_search(query: str) -> str:
    """
    사용자가 입력한 상품 설명(query)을 바탕으로 관련 HS 코드 정보와 품명 규격 가이드를 검색합니다.

    Args:
        query: 검색할 상품 설명 (예: "냉동 참치", "스마트워치", "노트북 컴퓨터")

    Returns:
        검색된 HS 코드 관련 정보 문자열 (예상 HS 코드, 품명, 관련 규정 포함)
    """
    from textwrap import shorten

    global _hs_code_search_call_count
    _hs_code_search_call_count += 1

    print(f"[Tool] hs_code_search 실행: call={_hs_code_search_call_count}, query={query}")

    # 하드 리밋: 에이전트 1회 실행당 최대 5회까지만 실제 검색 수행
    # (후보 3개 탐색을 위해 5회로 상향)
    if _hs_code_search_call_count > 5:
        return (
            "[검색 제한 초과] 더 이상 hs_code_search 도구를 호출할 수 없습니다. "
            "반드시 지금까지 수집한 정보를 바탕으로 최종 답변을 작성하세요. "
            "추가 도구 호출 없이 지정된 형식으로 즉시 응답하세요."
        )

    # 긴 문장은 핵심 키워드 2~3개로 자동 축약 (도메인 검색 팁 반영)
    # 예: "미국산 냉동 참치를 대형 유통용으로 수입하려고 합니다." -> "냉동 참치"
    raw = str(query)
    # 불필요한 조사/표현 일부 제거 (아주 가벼운 휴리스틱)
    for stop in ["수입하려고", "사용하는", "사용용", "위한", "용도", "대형", "소형"]:
        raw = raw.replace(stop, " ")

    tokens = raw.replace(",", " ").replace("(", " ").replace(")", " ").split()
    if len(tokens) > 3:
        # 한국어/한자/영문 혼합을 고려해 마지막 2~3 단어만 남김
        core_tokens = tokens[-3:]
        processed_query = " ".join(core_tokens)
    else:
        processed_query = raw.strip() or query

    # LLM 기반 검색 쿼리 확장
    expanded_query = expand_hs_search_query(processed_query)
    final_query = f"{processed_query} {expanded_query}".strip() if expanded_query else processed_query

    processed_query_short = shorten(final_query, width=120, placeholder="...")
    print(f"[Tool] hs_code_search 실제 검색어: {processed_query_short}")

    try:
        retrieved_docs = _balanced_search(final_query, k=5)
        if not retrieved_docs:
            return f"'{processed_query}'에 대한 HS 코드 정보를 찾을 수 없습니다."
        return "\n\n".join([doc.page_content for doc in retrieved_docs])
    except Exception as e:
        return f"HS 코드 검색 중 오류 발생: {str(e)}"


def parse_tariff_rate(rate_str: str) -> float:
    """'8%', '무세', NaN 같은 문자열 세율을 float 숫자로 변환"""
    if pd.isna(rate_str):
        return np.inf
    if isinstance(rate_str, str):
        if '무세' in rate_str:
            return 0.0
        try:
            return float(rate_str.replace('%', '').strip())
        except ValueError:
            return np.inf
    return float(rate_str)


@tool("tariff_search_by_hs_code")
def tariff_search_by_hs_code(hs_code: str) -> str:
    """
    정확한 HS 코드(세번)를 입력받아 관세율 정보를 조회합니다.
    기본세율, 잠정세율, WTO 협정세율을 비교하여 가장 낮은 세율을 최종 적용 세율로 결정합니다.
    
    Args:
        hs_code: HS 코드 (예: "0303.43-0000", "8517.62-9090")
    
    Returns:
        관세율 정보 문자열 (기본세율, 잠정세율, WTO 협정세율, 최종 적용 세율, 과세 단위)
    """
    global _tariff_search_call_count
    _tariff_search_call_count += 1

    print(f"[Tool] tariff_search_by_hs_code 실행: call={_tariff_search_call_count}, hs_code={hs_code}")

    # 하드 리밋: 에이전트 1회 실행당 최대 5회까지만 상세 관세 조회
    if _tariff_search_call_count > 5:
        return (
            "[조회 제한 초과] 더 이상 tariff_search_by_hs_code 도구를 호출할 수 없습니다. "
            "반드시 지금까지 수집한 정보를 바탕으로 최종 답변을 작성하세요. "
            "추가 도구 호출 없이 [결과] 형식으로 HS 코드와 관세율을 즉시 응답하세요."
        )

    if TARIFF_DF is None:
        return "오류: 관세율 정보 파일(tariff_by_hs.csv)이 로드되지 않았습니다."

    # 입력 HS 코드 정규화: 구분자 제거 + leading zeros 제거
    # CSV의 세번 컬럼은 앞자리 0이 없는 형태로 저장됨 (예: '0307494010' → '307494010')
    hs_code_cleaned = hs_code.replace('.', '').replace('-', '').lstrip('0')

    # CSV 세번도 leading zeros 제거 후 비교
    csv_codes = TARIFF_DF['세번'].astype(str).str.lstrip('0')

    # 1. 정확히 일치하는 행 우선 검색
    exact_mask = csv_codes == hs_code_cleaned
    if exact_mask.any():
        result_df = TARIFF_DF[exact_mask]
    else:
        # 2. 상위 자리수로 prefix 검색 (6자리 → 4자리 → 2자리 순으로 축소)
        result_df = pd.DataFrame()
        for prefix_len in (6, 4, 2):
            prefix = hs_code_cleaned[:prefix_len] if len(hs_code_cleaned) >= prefix_len else hs_code_cleaned
            prefix_mask = csv_codes.str.startswith(prefix)
            if prefix_mask.any():
                result_df = TARIFF_DF[prefix_mask]
                break

        if not result_df.empty:
            # 가장 세부적인 (세번 길이가 긴) 행만 남김 — 챕터 헤더 행 제외
            max_len = result_df['세번'].astype(str).str.len().max()
            result_df = result_df[result_df['세번'].astype(str).str.len() == max_len]

    if result_df.empty:
        return f"HS 코드 '{hs_code}'에 해당하는 관세율 정보를 찾을 수 없습니다."

    target_row = result_df.iloc[0]
    item_name = target_row.get('한글품명', '알 수 없음')

    # 세율 정보 추출 (실제 CSV 컬럼명 사용)
    basic_rate_str = target_row.get('기본세율 - A', '정보 없음')
    provisional_rate_str = target_row.get('잠정세율 - B', '정보 없음')
    wto_rate_str = target_row.get('WTO협정세율 - C', '정보 없음')

    # 단위 정보 추출
    weight_unit = target_row.get('중량단위', '정보 없음')
    quantity_unit = target_row.get('수량단위', '정보 없음')

    # 세율 변환 및 비교
    basic_rate = parse_tariff_rate(basic_rate_str)
    provisional_rate = parse_tariff_rate(provisional_rate_str)
    wto_rate = parse_tariff_rate(wto_rate_str)

    final_rate = min(basic_rate, provisional_rate, wto_rate)

    if final_rate == np.inf:
        return (
            f"HS 코드 '{hs_code}' ({item_name})에 대한 종가세(%) 정보를 찾을 수 없습니다. "
            f"종량세일 수 있으니 단위를 확인하세요. 중량단위: {weight_unit}, 수량단위: {quantity_unit}"
        )

    return (
        f"HS 코드 '{target_row['세번']}' ({item_name})의 세율 조회 결과:\n"
        f"- 기본세율: {basic_rate_str}\n"
        f"- 잠정세율: {provisional_rate_str}\n"
        f"- WTO 협정세율: {wto_rate_str}\n"
        f"-> 최종 적용 세율 (가장 낮은 값): {final_rate}%\n"
        f"- 과세 단위: 중량({weight_unit}), 수량({quantity_unit})"
    )


# ===== Tax Calculator 도구 =====

@tool("exchange_rate_loader")
def exchange_rate_loader(target_currency: str = "USD") -> Dict[str, Any]:
    """
    특정 국가의 통화(target_currency)와 대한민국 원(KRW) 사이의 현재 환율을 가져옵니다.
    
    Args:
        target_currency: 조회할 통화 코드 (기본값: USD, 예: EUR, JPY, CNY, KRW)
    
    Returns:
        환율 정보 딕셔너리 {"currency": str, "rate": float, "source": str}
    """
    print(f"[Tool] exchange_rate_loader 실행: currency={target_currency}")
    
    # KRW(원화)인 경우 환율 1.0 (변환 불필요)
    if target_currency.upper() == "KRW":
        return {
            "currency": "KRW",
            "rate": 1.0,
            "source": "KRW (no conversion needed)"
        }
    
    api_key = os.getenv("EXCHANGERATE_API_KEY")
    # 기본 환율 (API 없을 때 사용, 2024년 기준 대략적 값)
    default_rates = {
        "USD": 1350.0,   # 미국 달러
        "EUR": 1450.0,   # 유로
        "JPY": 9.0,      # 일본 엔
        "CNY": 185.0,    # 중국 위안
        "GBP": 1700.0,   # 영국 파운드
        "AUD": 880.0,    # 호주 달러
        "CAD": 1000.0,   # 캐나다 달러
        "CHF": 1500.0,   # 스위스 프랑
        "HKD": 175.0,    # 홍콩 달러
        "SGD": 1000.0,   # 싱가포르 달러
        "TWD": 42.0,     # 대만 달러
        "THB": 38.0,     # 태국 바트
        "VND": 0.055,    # 베트남 동
        "INR": 16.0,     # 인도 루피
        "MYR": 285.0,    # 말레이시아 링깃
        "PHP": 24.0,     # 필리핀 페소
        "IDR": 0.085,    # 인도네시아 루피아
        "RUB": 15.0,     # 러시아 루블
        "BRL": 270.0,    # 브라질 레알
        "MXN": 78.0,     # 멕시코 페소
        "NZD": 820.0,    # 뉴질랜드 달러
        "SEK": 125.0,    # 스웨덴 크로나
        "NOK": 125.0,    # 노르웨이 크로네
        "DKK": 195.0,    # 덴마크 크로네
        "AED": 370.0,    # UAE 디르함
        "SAR": 360.0,    # 사우디 리얄
    }
    
    if not api_key:
        rate = default_rates.get(target_currency.upper(), 1350.0)
        return {
            "currency": target_currency.upper(),
            "rate": rate,
            "source": "default (API key not found)"
        }
    
    url = f"https://v6.exchangerate-api.com/v6/{api_key}/latest/{target_currency}"
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            data = response.json()
            rate = data['conversion_rates'].get('KRW', default_rates.get(target_currency.upper(), 1350.0))
            return {
                "currency": target_currency.upper(),
                "rate": rate,
                "source": "exchangerate-api.com"
            }
    except requests.RequestException as e:
        print(f"[Tool] 환율 API 오류: {e}")
    
    rate = default_rates.get(target_currency.upper(), 1350.0)
    return {
        "currency": target_currency.upper(),
        "rate": rate,
        "source": "default (API error)"
    }


@tool("final_cost_calculator")
def final_cost_calculator(
    item_price: float, 
    quantity: int, 
    exchange_rate: float, 
    tariff_rate: float
) -> Dict[str, Any]:
    """
    상품 단가, 수량, 환율, 관세율을 입력받아 총 예상 수입 원가를 계산합니다.
    부가세(10%)는 관세가 포함된 금액에 부과됩니다.
    
    Args:
        item_price: 상품 단가 (외화)
        quantity: 수량
        exchange_rate: 환율 (KRW 기준)
        tariff_rate: 관세율 (%, 예: 8은 8%를 의미)
    
    Returns:
        계산 결과 딕셔너리 (각 단계별 금액 및 최종 비용)
    """
    print(f"[Tool] final_cost_calculator 실행: price={item_price}, qty={quantity}, rate={exchange_rate}, tariff={tariff_rate}")
    
    # 계산
    total_item_price_foreign = item_price * quantity
    total_item_price_krw = total_item_price_foreign * exchange_rate
    tariff = total_item_price_krw * (tariff_rate / 100)
    price_with_tariff = total_item_price_krw + tariff
    vat = price_with_tariff * 0.10
    total_cost = price_with_tariff + vat
    
    return {
        "total_item_price_foreign": total_item_price_foreign,
        "total_item_price_krw": total_item_price_krw,
        "tariff_rate_percent": tariff_rate,
        "tariff_amount": tariff,
        "vat_amount": vat,
        "total_cost": total_cost,
        "breakdown": (
            f"--- 최종 비용 계산 결과 ---\n"
            f"1. 총 물품 가격 (원화): {total_item_price_krw:,.0f} 원\n"
            f"   (단가 {item_price:,.2f} × 수량 {quantity} × 환율 {exchange_rate:,.2f})\n"
            f"2. 예상 관세 ({tariff_rate}%): {tariff:,.0f} 원\n"
            f"3. 예상 부가세 (10%): {vat:,.0f} 원\n"
            f"--------------------------------\n"
            f"   총 예상 수입 비용: {total_cost:,.0f} 원\n"
            f"--------------------------------"
        )
    }


# ===== Report Writer 도구 =====

@tool("pdf_report_exporter")
def pdf_report_exporter(report_content: str, filename: str = "report.pdf") -> str:
    """
    분석 결과를 담은 문자열(report_content)을 전문적인 디자인의 PDF 파일로 저장합니다.
    
    Args:
        report_content: 보고서 내용 문자열
        filename: 저장할 파일명 (기본값: report.pdf)
    
    Returns:
        저장 결과 메시지
    """
    print(f"[Tool] pdf_report_exporter 실행: filename={filename}")
    
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib.colors import HexColor, black, white
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        
        # 한글 폰트 등록
        font_name = 'Helvetica'
        font_paths = [
            './fonts/NanumGothic.ttf',
            'fonts/NanumGothic.ttf',
            '/usr/share/fonts/truetype/nanum/NanumGothic.ttf',
            'NanumGothic.ttf',
            os.path.join(os.path.dirname(__file__), '..', 'fonts', 'NanumGothic.ttf'),
        ]
        
        for font_path in font_paths:
            try:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('NanumGothic', font_path))
                    font_name = 'NanumGothic'
                    print(f"[Tool] 폰트 로드 성공: {font_path}")
                    break
            except Exception:
                continue
        
        # 문서 생성
        doc = SimpleDocTemplate(
            filename,
            pagesize=A4,
            rightMargin=20*mm,
            leftMargin=20*mm,
            topMargin=20*mm,
            bottomMargin=20*mm
        )
        
        # 스타일 정의
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=font_name,
            fontSize=14,
            textColor=HexColor('#1a365d'),
            alignment=TA_CENTER,
            spaceAfter=12,
            spaceBefore=0,
        )
        
        section_style = ParagraphStyle(
            'CustomSection',
            parent=styles['Heading2'],
            fontName=font_name,
            fontSize=10,
            textColor=HexColor('#2c5282'),
            spaceBefore=10,
            spaceAfter=6,
            borderColor=HexColor('#e2e8f0'),
            borderWidth=0,
            borderPadding=4,
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=9,
            textColor=HexColor('#2d3748'),
            spaceBefore=2,
            spaceAfter=2,
            leading=14,
        )
        
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=9,
            textColor=HexColor('#4a5568'),
            leftIndent=15,
            spaceBefore=1,
            spaceAfter=1,
            leading=13,
        )
        
        # 콘텐츠 빌드
        story = []
        
        # 헤더 라인
        story.append(HRFlowable(width="100%", thickness=2, color=HexColor('#3182ce'), spaceBefore=0, spaceAfter=10))
        
        # 제목
        story.append(Paragraph("HS 코드 및 수입 원가 분석 보고서", title_style))
        
        # 날짜
        from datetime import datetime
        date_style = ParagraphStyle('Date', parent=body_style, alignment=TA_CENTER, textColor=HexColor('#718096'), fontSize=8)
        story.append(Paragraph(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}", date_style))
        
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor('#e2e8f0'), spaceBefore=10, spaceAfter=15))
        
        # 본문 파싱
        for line in report_content.split('\n'):
            s = line.strip()
            if not s:
                story.append(Spacer(1, 4))
                continue
            
            # 마크다운 정리
            clean = s.replace('**', '').replace('*', '').replace('>', '').strip()
            if not clean:
                continue
            
            # 섹션 제목 (## 또는 #)
            if s.startswith('## ') or s.startswith('# '):
                section_text = clean.lstrip('# ').strip()
                story.append(Spacer(1, 6))
                story.append(HRFlowable(width="30%", thickness=1, color=HexColor('#3182ce'), spaceBefore=0, spaceAfter=4))
                story.append(Paragraph(section_text, section_style))
            # 불릿 리스트
            elif s.startswith('- ') or s.startswith('* '):
                item_text = clean.lstrip('-* ').strip()
                story.append(Paragraph(f"• {item_text}", bullet_style))
            # 일반 본문
            else:
                story.append(Paragraph(clean, body_style))
        
        # 푸터
        story.append(Spacer(1, 20))
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor('#e2e8f0'), spaceBefore=10, spaceAfter=6))
        footer_style = ParagraphStyle('Footer', parent=body_style, alignment=TA_CENTER, textColor=HexColor('#a0aec0'), fontSize=7)
        story.append(Paragraph("본 보고서는 TradeSimple AI에 의해 자동 생성되었습니다. 실제 수입 시 관세사 확인을 권장합니다.", footer_style))
        
        # PDF 생성
        doc.build(story)
        
        return f"PDF 보고서가 '{filename}'로 성공적으로 저장되었습니다."
    except Exception as e:
        return f"PDF 저장 중 오류 발생: {str(e)}"


@tool("word_report_exporter")
def word_report_exporter(report_content: str, filename: str = "report.docx") -> str:
    """
    분석 결과를 담은 문자열(report_content)을 전문적인 디자인의 Word 파일로 저장합니다.
    파란 헤더 배너, 섹션별 배경색, 인라인 볼드 처리, 푸터를 포함합니다.

    Args:
        report_content: 보고서 내용 문자열 (마크다운)
        filename: 저장할 파일명 (기본값: report.docx)

    Returns:
        저장 결과 메시지
    """
    print(f"[Tool] word_report_exporter 실행: filename={filename}")

    try:
        import re as _re
        from datetime import datetime as _dt
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        FONT = "Nanum Gothic"
        BLUE_DARK_HEX  = "1A366B"   # header bg
        BLUE_MID       = RGBColor(0x2B, 0x5B, 0xA6)
        BLUE_MID_HEX   = "2B5BA6"
        BLUE_LIGHT_HEX = "EBF2FB"   # section header bg
        GRAY_LIGHT_HEX = "F7F7F7"   # footer bg
        WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
        GRAY_TEXT      = RGBColor(0x88, 0x88, 0x88)

        doc = DocxDocument()

        # ── 문서 여백 ───────────────────────────────────────────────────
        for sec in doc.sections:
            sec.left_margin   = Inches(1.0)
            sec.right_margin  = Inches(1.0)
            sec.top_margin    = Inches(0.8)
            sec.bottom_margin = Inches(0.8)

        # ── XML 헬퍼 ───────────────────────────────────────────────────
        def _set_cell_shd(cell, fill_hex):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill_hex)
            tcPr.append(shd)

        def _set_para_shd(para, fill_hex):
            pPr = para._element.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill_hex)
            pPr.append(shd)

        def _set_para_left_border(para, color_hex=BLUE_MID_HEX, sz="28"):
            pPr = para._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), sz)
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), color_hex)
            pBdr.append(left)
            pPr.append(pBdr)

        def _remove_table_borders(table):
            tbl_el = table._element
            tblPr = tbl_el.find(qn("w:tblPr"))
            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")
                tbl_el.insert(0, tblPr)
            old = tblPr.find(qn("w:tblBorders"))
            if old is not None:
                tblPr.remove(old)
            tblBorders = OxmlElement("w:tblBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "none")
                b.set(qn("w:sz"), "0")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), "auto")
                tblBorders.append(b)
            tblPr.append(tblBorders)

        def _add_run(para, text, bold=False, size=10, color=None, italic=False):
            """단락에 런 추가 (한글 폰트 포함)."""
            run = para.add_run(text)
            run.font.name = FONT
            run.font.size = Pt(size)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = color
            # 동아시아 폰트 명시 설정
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), FONT)
            rFonts.set(qn("w:hAnsi"), FONT)
            rFonts.set(qn("w:eastAsia"), FONT)
            rFonts.set(qn("w:cs"), FONT)
            return run

        def _parse_inline(para, text, base_size=9, base_bold=False, base_color=None):
            """**bold** 인라인 마커를 처리하여 런을 분리 추가."""
            parts = _re.split(r"\*\*(.+?)\*\*", text)
            for i, part in enumerate(parts):
                if not part:
                    continue
                is_bold = (i % 2 == 1) or base_bold
                _add_run(para, part, bold=is_bold, size=base_size, color=base_color)

        # ── 파란 헤더 배너 (1×1 표로 구현) ─────────────────────────────
        header_tbl = doc.add_table(rows=1, cols=1)
        header_tbl.style = "Table Grid"
        _remove_table_borders(header_tbl)

        hcell = header_tbl.cell(0, 0)
        _set_cell_shd(hcell, BLUE_DARK_HEX)

        # 셀 내부 여백 (상하 120, 좌우 200 DXA)
        tc = hcell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for side, val in (("top", "120"), ("bottom", "120"), ("left", "200"), ("right", "200")):
            m = OxmlElement(f"w:{side}")
            m.set(qn("w:w"), val)
            m.set(qn("w:type"), "dxa")
            tcMar.append(m)
        tcPr.append(tcMar)

        # 헤더 제목
        hp = hcell.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.space_before = Pt(8)
        hp.paragraph_format.space_after  = Pt(2)
        _add_run(hp, "HS 코드 및 수입 원가 분석 보고서", bold=True, size=15, color=WHITE)

        # 헤더 부제
        sp = hcell.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(8)
        _add_run(
            sp,
            f"TradeSimple AI  |  작성일: {_dt.now().strftime('%Y년 %m월 %d일')}",
            bold=False, size=9,
            color=RGBColor(0xCC, 0xD9, 0xF5),
        )

        # 헤더 아래 여백
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # ── 본문: 마크다운 라인 파싱 ────────────────────────────────────
        for line in report_content.split("\n"):
            s = line.strip()

            # 빈 줄
            if not s:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(2)
                continue

            # H1 / H2 섹션 헤더
            if s.startswith("## ") or s.startswith("# "):
                title_text = _re.sub(r"^#+\s*", "", s).strip()
                title_text = _re.sub(r"\*\*(.+?)\*\*", r"\1", title_text)

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after  = Pt(5)
                p.paragraph_format.left_indent  = Cm(0.5)
                _set_para_shd(p, BLUE_LIGHT_HEX)
                _set_para_left_border(p)
                _add_run(p, title_text, bold=True, size=11, color=BLUE_MID)

            # 불릿 리스트
            elif s.startswith("- ") or s.startswith("* "):
                item_text = _re.sub(r"^[-*]\s+", "", s).strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent  = Cm(0.9)
                p.paragraph_format.space_after  = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                _add_run(p, "▸  ", bold=False, size=9, color=BLUE_MID)
                _parse_inline(p, item_text, base_size=9)

            # 일반 본문 (인용 블록 > 포함)
            else:
                clean = _re.sub(r"^>+\s*", "", s).strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                _parse_inline(p, clean, base_size=9)

        # ── 푸터 ────────────────────────────────────────────────────────
        doc.add_paragraph().paragraph_format.space_after = Pt(8)
        fp = doc.add_paragraph()
        _set_para_shd(fp, GRAY_LIGHT_HEX)
        fp.paragraph_format.space_before = Pt(6)
        fp.paragraph_format.space_after  = Pt(6)
        fp.paragraph_format.left_indent  = Cm(0.3)
        _add_run(
            fp,
            "※ 본 보고서는 TradeSimple AI에 의해 자동 생성되었습니다. "
            "실제 수입 시 관세사 또는 세관의 확인을 권장합니다.",
            bold=False, size=8, color=GRAY_TEXT,
        )

        doc.save(filename)
        return f"Word 보고서가 '{filename}'로 성공적으로 저장되었습니다."

    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"Word 저장 중 오류 발생: {str(e)}"


@tool("excel_report_exporter")
def excel_report_exporter(data: Dict[str, Any], filename: str = "report.xlsx") -> str:
    """
    분석 결과 데이터를 Excel 파일로 저장합니다.
    
    Args:
        data: 보고서 데이터 딕셔너리 (예: {"물품명": "스마트워치", "HS코드": "8517.62", ...})
        filename: 저장할 파일명 (기본값: report.xlsx)
    
    Returns:
        저장 결과 메시지
    """
    print(f"[Tool] excel_report_exporter 실행: filename={filename}")
    
    try:
        # 딕셔너리를 DataFrame으로 변환
        if isinstance(data, dict):
            df = pd.DataFrame([data])
        elif isinstance(data, list):
            df = pd.DataFrame(data)
        else:
            return "엑셀 파일 생성 실패: 입력 데이터는 딕셔너리 또는 리스트여야 합니다."
        
        df.to_excel(filename, index=False)
        return f"Excel 보고서가 '{filename}'로 성공적으로 저장되었습니다."
    except Exception as e:
        return f"Excel 저장 중 오류 발생: {str(e)}"


# ===== 도구 그룹 정의 =====

# HS Code & Tax Finder 에이전트용 도구
hs_code_finder_tools = [hs_code_search, tariff_search_by_hs_code]

# Tax Calculator 에이전트용 도구
tax_calculator_tools = [exchange_rate_loader, final_cost_calculator]

# Report Writer 에이전트용 도구
report_writer_tools = [pdf_report_exporter, word_report_exporter, excel_report_exporter]

# 모든 도구 리스트
all_tools = hs_code_finder_tools + tax_calculator_tools + report_writer_tools
